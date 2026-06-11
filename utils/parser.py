"""
Resume Parser Module

Handles text extraction from PDF and DOCX resume files.
Uses pdfplumber as primary PDF extractor with PyPDF2 fallback.
Uses python-docx for DOCX files.
"""

import re
import io
from typing import Optional

import pdfplumber
from PyPDF2 import PdfReader
from docx import Document


def extract_text_from_pdf(file) -> str:
    """
    Extract text from a PDF file using pdfplumber (primary) with PyPDF2 fallback.

    Args:
        file: A file-like object (e.g., Streamlit UploadedFile).

    Returns:
        Extracted text as a string.
    """
    text = ""

    # Primary: pdfplumber (better layout preservation)
    try:
        file.seek(0)
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception:
        text = ""

    # Fallback: PyPDF2 if pdfplumber returned nothing
    if not text.strip():
        try:
            file.seek(0)
            reader = PdfReader(file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception:
            text = ""

    return text.strip()


def extract_text_from_docx(file) -> str:
    """
    Extract text from a DOCX file using python-docx.

    Extracts content from paragraphs and tables.

    Args:
        file: A file-like object (e.g., Streamlit UploadedFile).

    Returns:
        Extracted text as a string.
    """
    text = ""
    try:
        file.seek(0)
        doc = Document(file)

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Check if paragraph is a list item
                style_name = paragraph.style.name if paragraph.style else ""
                if "List" in style_name or "Bullet" in style_name:
                    text += "- " + paragraph.text + "\n"
                else:
                    text += paragraph.text + "\n"

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text += row_text + "\n"
    except Exception:
        text = ""

    return text.strip()


def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text.

    Removes extra whitespace, special characters, and normalizes formatting.

    Args:
        text: Raw extracted text.

    Returns:
        Cleaned and normalized text.
    """
    if not text:
        return ""

    # Replace common unicode characters
    text = text.replace("\u2019", "'")
    text = text.replace("\u2018", "'")
    text = text.replace("\u201c", '"')
    text = text.replace("\u201d", '"')
    text = text.replace("\u2013", "-")
    text = text.replace("\u2014", "-")
    text = text.replace("\u2022", " ")
    text = text.replace("\u25cf", " ")
    text = text.replace("\u25cb", " ")
    text = text.replace("\u25a0", " ")
    text = text.replace("\u25aa", " ")

    # Remove non-printable characters (keep newlines and tabs)
    text = re.sub(r'[^\x20-\x7E\n\t]', ' ', text)

    # Normalize whitespace
    text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs → single space
    text = re.sub(r'\n\s*\n', '\n\n', text)  # Multiple blank lines → double newline
    text = re.sub(r'\n{3,}', '\n\n', text)  # Cap at 2 newlines

    # Strip leading/trailing whitespace per line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)

    return text.strip()


def extract_text(file) -> dict:
    """
    Main entry point for text extraction.

    Detects file type and dispatches to the appropriate extractor.

    Args:
        file: A Streamlit UploadedFile object.

    Returns:
        Dictionary with keys:
        - filename (str): Name of the uploaded file
        - text (str): Extracted and cleaned text
        - file_type (str): 'pdf', 'docx', or 'unknown'
        - success (bool): Whether extraction succeeded
        - error (str | None): Error message if extraction failed
    """
    filename = getattr(file, 'name', 'unknown')
    result = {
        "filename": filename,
        "text": "",
        "file_type": "unknown",
        "success": False,
        "error": None,
    }

    # Determine file type
    lower_name = filename.lower()
    if lower_name.endswith('.pdf'):
        result["file_type"] = "pdf"
        raw_text = extract_text_from_pdf(file)
    elif lower_name.endswith('.docx'):
        result["file_type"] = "docx"
        raw_text = extract_text_from_docx(file)
    elif lower_name.endswith('.doc'):
        result["file_type"] = "doc"
        result["error"] = "Legacy .doc format is not supported. Please convert to .docx or .pdf."
        return result
    else:
        result["error"] = f"Unsupported file format: {lower_name.split('.')[-1]}"
        return result

    # Validate extracted text
    if not raw_text or len(raw_text.strip()) < 50:
        result["error"] = (
            "Could not extract sufficient text from this file. "
            "The file may be empty, scanned (image-based), or corrupted."
        )
        # Still include whatever text was found
        result["text"] = clean_text(raw_text) if raw_text else ""
        return result

    result["text"] = clean_text(raw_text)
    result["success"] = True
    return result
